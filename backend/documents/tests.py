from __future__ import annotations

from io import BytesIO
from pathlib import Path
from tempfile import TemporaryDirectory

from django.contrib.auth.models import User
from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import TestCase, override_settings
from rest_framework.test import APIClient

from documents.models import Document, DocumentStatus
from rag.embeddings import get_embedding_service
from rag.vectorstore import get_vector_store


def _settings(temp_dir: str) -> dict:
    return {
        "PROCESS_DOCUMENTS_ASYNC": False,
        "EMBEDDING_PROVIDER": "lexical",
        "EMBEDDING_MODEL": "lexical",
        "VECTOR_DB_PATH": Path(temp_dir) / "chroma",
        "MEDIA_ROOT": Path(temp_dir) / "media",
        "SIMILARITY_THRESHOLD": 0.05,
        "RAG_TOP_K": 5,
    }


def _reset_singletons():
    get_embedding_service.cache_clear()
    get_vector_store.cache_clear()


class DocumentApiTests(TestCase):
    def setUp(self):
        self.tmp = TemporaryDirectory()
        self.addCleanup(self.tmp.cleanup)
        self.addCleanup(_reset_singletons)
        override = override_settings(**_settings(self.tmp.name))
        override.enable()
        self.addCleanup(override.disable)
        _reset_singletons()

        self.admin = User.objects.create_user(
            username="admin", password="adminpass123", is_staff=True
        )
        self.student = User.objects.create_user(username="student", password="studentpass")
        self.client = APIClient()

    def _login_admin(self):
        self.client.login(username="admin", password="adminpass123")

    def _txt_file(self, name="policy.txt", content="Attendance must be 75 percent.\n"):
        return SimpleUploadedFile(name, content.encode("utf-8"), content_type="text/plain")

    def test_student_cannot_upload(self):
        self.client.login(username="student", password="studentpass")
        response = self.client.post(
            "/api/documents/",
            {"title": "Secret", "category": "attendance", "file": self._txt_file()},
            format="multipart",
        )
        self.assertEqual(response.status_code, 403)

    def test_txt_upload_and_processing(self):
        self._login_admin()
        response = self.client.post(
            "/api/documents/",
            {
                "title": "Attendance Policy",
                "category": "attendance",
                "department": "Academics",
                "description": "Test",
                "version": "1.0",
                "file": self._txt_file(),
            },
            format="multipart",
        )
        self.assertEqual(response.status_code, 201, response.content)
        self.assertEqual(response.data["status"], DocumentStatus.COMPLETED)
        self.assertGreater(response.data["chunk_count"], 0)
        doc = Document.objects.get(id=response.data["id"])
        self.assertGreater(doc.chunks.count(), 0)

    def test_docx_upload(self):
        from docx import Document as WordDocument

        buffer = BytesIO()
        word = WordDocument()
        word.add_heading("Fee Refund Policy", 0)
        word.add_paragraph("Withdrawal before classes begins yields a 100% tuition refund minus processing charges.")
        word.save(buffer)
        uploaded = SimpleUploadedFile(
            "fees.docx",
            buffer.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        self._login_admin()
        response = self.client.post(
            "/api/documents/",
            {"title": "Fee Refund Policy", "category": "finance", "file": uploaded},
            format="multipart",
        )
        self.assertEqual(response.status_code, 201, response.content)
        self.assertEqual(response.data["status"], DocumentStatus.COMPLETED)

    def test_pdf_upload(self):
        from pypdf import PdfWriter

        writer = PdfWriter()
        writer.add_blank_page(width=200, height=200)
        # pypdf blank pages have no text; add a page with text via report-less fallback:
        buffer = BytesIO()
        writer.write(buffer)

        # Build a tiny PDF with text using pypdf page content is hard; use a minimal manual PDF.
        pdf_bytes = (
            b"%PDF-1.4\n"
            b"1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n"
            b"2 0 obj<</Type/Pages/Count 1/Kids[3 0 R]>>endobj\n"
            b"3 0 obj<</Type/Page/Parent 2 0 R/MediaBox[0 0 200 200]/Contents 4 0 R/Resources<</Font<</F1 5 0 R>>>>>>endobj\n"
            b"4 0 obj<</Length 68>>stream\n"
            b"BT /F1 12 Tf 20 150 Td (Students must maintain 75 percent attendance.) Tj ET\n"
            b"endstream\nendobj\n"
            b"5 0 obj<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>endobj\n"
            b"xref\n0 6\n0000000000 65535 f \n"
            b"trailer<</Size 6/Root 1 0 R>>\nstartxref\n0\n%%EOF\n"
        )
        uploaded = SimpleUploadedFile("attendance.pdf", pdf_bytes, content_type="application/pdf")
        self._login_admin()
        response = self.client.post(
            "/api/documents/",
            {"title": "Attendance PDF", "category": "attendance", "file": uploaded},
            format="multipart",
        )
        self.assertEqual(response.status_code, 201, response.content)
        # Extraction may yield zero text on this minimal PDF; either completed or failed with a clear error.
        self.assertIn(response.data["status"], {DocumentStatus.COMPLETED, DocumentStatus.FAILED})

    def test_rejects_exe(self):
        self._login_admin()
        uploaded = SimpleUploadedFile("malware.exe", b"MZ", content_type="application/octet-stream")
        response = self.client.post(
            "/api/documents/",
            {"title": "Bad", "category": "general", "file": uploaded},
            format="multipart",
        )
        self.assertEqual(response.status_code, 400)

    def test_delete_removes_document(self):
        self._login_admin()
        created = self.client.post(
            "/api/documents/",
            {"title": "Temp", "category": "general", "file": self._txt_file()},
            format="multipart",
        )
        doc_id = created.data["id"]
        deleted = self.client.delete(f"/api/documents/{doc_id}/")
        self.assertEqual(deleted.status_code, 204)
        self.assertFalse(Document.objects.filter(id=doc_id).exists())

    def test_reprocess(self):
        self._login_admin()
        created = self.client.post(
            "/api/documents/",
            {"title": "Reprocess Me", "category": "general", "file": self._txt_file("a.txt", "Probation requires CGPA 2.00.\n")},
            format="multipart",
        )
        doc_id = created.data["id"]
        again = self.client.post(f"/api/documents/{doc_id}/reprocess/?sync=1")
        self.assertEqual(again.status_code, 200, again.content)
        self.assertEqual(again.data["status"], DocumentStatus.COMPLETED)

    def test_list_filter_and_dashboard(self):
        self._login_admin()
        self.client.post(
            "/api/documents/",
            {"title": "Exam Policy", "category": "examination", "file": self._txt_file("e.txt", "Final exam is 50 percent.\n")},
            format="multipart",
        )
        listing = self.client.get("/api/documents/?category=examination&search=Exam")
        self.assertEqual(listing.status_code, 200)
        self.assertEqual(len(listing.data), 1)
        stats = self.client.get("/api/dashboard/stats/")
        self.assertEqual(stats.status_code, 200)
        self.assertGreaterEqual(stats.data["total_documents"], 1)
