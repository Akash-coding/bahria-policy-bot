from __future__ import annotations

import logging
import re
from pathlib import Path

logger = logging.getLogger("rag")


class ExtractionError(RuntimeError):
    pass


def _clean_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r" *\n *", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_pages(file_path: str | Path, file_type: str) -> list[tuple[int | None, str]]:
    path = Path(file_path)
    if not path.exists():
        raise ExtractionError(f"File not found: {path}")

    file_type = (file_type or path.suffix.lstrip(".")).lower()
    if file_type == "pdf":
        return _extract_pdf(path)
    if file_type == "docx":
        return _extract_docx(path)
    if file_type == "txt":
        return _extract_txt(path)
    raise ExtractionError(f"Unsupported file type: {file_type}")


def _extract_pdf(path: Path) -> list[tuple[int | None, str]]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ExtractionError("pypdf is required to process PDF files.") from exc

    try:
        reader = PdfReader(str(path))
    except Exception as exc:
        raise ExtractionError(f"Unable to read PDF: {exc}") from exc

    pages: list[tuple[int | None, str]] = []
    for index, page in enumerate(reader.pages, start=1):
        try:
            raw = page.extract_text() or ""
        except Exception:
            logger.exception("Failed to extract text from PDF page %s of %s", index, path)
            raw = ""
        cleaned = _clean_text(raw)
        if cleaned:
            pages.append((index, cleaned))
    if not pages:
        raise ExtractionError("No extractable text was found in the PDF.")
    return pages


def _extract_docx(path: Path) -> list[tuple[int | None, str]]:
    try:
        import docx
    except ImportError as exc:
        raise ExtractionError("python-docx is required to process DOCX files.") from exc

    try:
        document = docx.Document(str(path))
    except Exception as exc:
        raise ExtractionError(f"Unable to read DOCX: {exc}") from exc

    paragraphs = [_clean_text(p.text) for p in document.paragraphs]
    tables_text: list[str] = []
    for table in document.tables:
        rows = []
        for row in table.rows:
            cells = [_clean_text(cell.text) for cell in row.cells]
            rows.append(" | ".join(cell for cell in cells if cell))
        tables_text.append("\n".join(row for row in rows if row))

    body = _clean_text("\n\n".join([p for p in paragraphs if p] + tables_text))
    if not body:
        raise ExtractionError("No extractable text was found in the Word document.")
    return [(None, body)]


def _extract_txt(path: Path) -> list[tuple[int | None, str]]:
    try:
        text = path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        text = path.read_text(encoding="utf-8", errors="replace")
    cleaned = _clean_text(text)
    if not cleaned:
        raise ExtractionError("The text file is empty.")
    return [(1, cleaned)]
